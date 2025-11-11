from __future__ import annotations

import json
from http import HTTPStatus
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import app
from app.services.templates import (
    TemplateCreationError,
    TemplateCreationResult,
    create_template_from_upload,
    template_repository,
)


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Dear {{client_name}},")
    document.add_paragraph("The agreement will start on {{sign_date}}.")
    document.add_paragraph("Regards, {{client_name}} team.")
    document.save(path)


def test_create_template_from_upload(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source)

    result: TemplateCreationResult = create_template_from_upload(
        file_name="Client Agreement.docx",
        file_bytes=source.read_bytes(),
        template_root=tmp_path,
    )

    assert result.metadata.id == "client-agreement"
    assert result.metadata.entry == "Client Agreement.docx"
    assert result.metadata.description == ""
    assert [field.name for field in result.metadata.fields] == ["client_name", "sign_date"]

    metadata_payload = json.loads(result.metadata_path.read_text(encoding="utf-8"))
    assert metadata_payload["id"] == "client-agreement"
    assert metadata_payload["entry"] == "Client Agreement.docx"

    # Creating another template with the same name should append a numeric suffix.
    result_second = create_template_from_upload(
        file_name="Client Agreement.docx",
        file_bytes=source.read_bytes(),
        template_root=tmp_path,
    )
    assert result_second.metadata.id == "client-agreement-1"


def test_create_template_from_upload_rejects_invalid_extension(tmp_path: Path) -> None:
    with pytest.raises(TemplateCreationError):
        create_template_from_upload(
            file_name="notes.txt",
            file_bytes=b"hello",
            template_root=tmp_path,
        )


def test_upload_template_endpoint(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = tmp_path / "upload.docx"
    _write_docx(source)

    monkeypatch.setattr(settings, "template_root_relative", tmp_path)
    monkeypatch.setattr(template_repository, "template_root", settings.template_root)
    template_repository.refresh()

    client = TestClient(app)

    with source.open("rb") as fp:
        response = client.post(
            "/api/templates/upload",
            files={
                "file": (
                    "Partner Contract.docx",
                    fp.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == HTTPStatus.CREATED
    body = response.json()

    assert body["template"]["id"] == "partner-contract"
    assert body["metadata_path"] == "partner-contract/metadata.json"
    assert "metadata.json" in body["message"]

    metadata_file = tmp_path / body["metadata_path"]
    assert metadata_file.exists()
    saved = json.loads(metadata_file.read_text(encoding="utf-8"))
    assert saved["fields"] == [
        {"name": "client_name", "type": "string"},
        {"name": "sign_date", "type": "string"},
    ]

    # Repository should now expose the newly created template.
    template = template_repository.get_template(body["template"]["id"])
    assert template.entry == "Partner Contract.docx"
