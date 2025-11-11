from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from app.models.templates import TemplateMetadata
from app.services.docx_renderer import DocxRenderService


@pytest.fixture()
def template_metadata() -> TemplateMetadata:
    return TemplateMetadata(
        id="test_template",
        name="测试模板",
        entry="template.docx",
        fields=[],
    )


@pytest.fixture()
def template_root(tmp_path: Path, template_metadata: TemplateMetadata) -> Path:
    template_dir = tmp_path / template_metadata.id
    template_dir.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_paragraph("甲方：{{ party_a_name }}")
    document.add_paragraph("乙方：{{ party_b_name }}")
    document.add_paragraph("签署日期：{{ sign_date }}")
    document.add_paragraph("Logo：{{ logo }}")
    document.save(template_dir / template_metadata.entry)

    image_path = template_dir / "logo.png"
    Image.new("RGB", (10, 10), color="red").save(image_path)

    return tmp_path


def test_docx_renderer_supports_basic_fields_and_images(
    template_root: Path,
    template_metadata: TemplateMetadata,
    tmp_path: Path,
) -> None:
    renderer = DocxRenderService(template_root=template_root)
    output_path = tmp_path / "output.docx"

    renderer.render(
        template_metadata,
        {
            "party_a_name": "甲方公司",
            "party_b_name": "乙方公司",
            "sign_date": date(2025, 1, 1),
            "logo": {
                "__type": "image",
                "path": template_root / template_metadata.id / "logo.png",
                "width_mm": 15,
            },
        },
        output_path=output_path,
    )

    assert output_path.exists()

    rendered = Document(output_path)
    full_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)

    assert "甲方：甲方公司" in full_text
    assert "乙方：乙方公司" in full_text
    assert "签署日期：2025-01-01" in full_text
    assert rendered.inline_shapes, "图像未写入文档"


