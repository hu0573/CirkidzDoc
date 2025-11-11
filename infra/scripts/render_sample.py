#!/usr/bin/env python3
import json
import os
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate


WORKSPACE = Path("/workspace")
TMP_DIR = Path(os.environ.get("RENDER_TMP_DIR", WORKSPACE / "tmp"))
OUTPUT_DIR = Path(os.environ.get("RENDER_OUTPUT_DIR", WORKSPACE / "output"))


def ensure_directories() -> None:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def build_template(template_path: Path) -> None:
    if template_path.exists():
        return
    doc = Document()
    doc.add_heading("示例服务协议", level=1)
    doc.add_paragraph("甲方：{{ party_a }}")
    doc.add_paragraph("乙方：{{ party_b }}")
    doc.add_paragraph("签署日期：{{ sign_date }}")
    doc.add_paragraph("")
    doc.add_paragraph("主要条款摘要：")
    doc.add_paragraph("1. 服务范围：{{ clause_1 }}")
    doc.add_paragraph("2. 费用条款：{{ clause_2 }}")
    doc.add_paragraph("3. 联系方式：{{ contact }}")
    doc.add_paragraph("")
    doc.add_paragraph("（以上内容由自动化渲染流程填充生成）")
    doc.save(template_path)


def render_docx(template_path: Path, output_path: Path) -> dict:
    context = {
        "party_a": "奇客科技有限公司",
        "party_b": "示例客户",
        "sign_date": date.today().isoformat(),
        "clause_1": "提供文档模板自动填充与格式转换服务。",
        "clause_2": "按实际使用量结算，试点阶段免服务费。",
        "contact": "support@example.com",
    }
    template = DocxTemplate(str(template_path))
    template.render(context)
    template.save(output_path)
    return context


def convert_with_libreoffice(docx_path: Path, output_dir: Path) -> Path:
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf:writer_pdf_Export",
            str(docx_path),
            "--outdir",
            str(output_dir),
        ],
        check=True,
    )
    return output_dir / f"{docx_path.stem}.pdf"


def convert_with_pandoc(docx_path: Path, output_dir: Path, fmt: str) -> Path:
    target = output_dir / f"{docx_path.stem}.{fmt}"
    subprocess.run(
        [
            "pandoc",
            str(docx_path),
            "-o",
            str(target),
        ],
        check=True,
    )
    return target


def summarize(results: dict) -> None:
    print(json.dumps(results, indent=2, ensure_ascii=False))


def main() -> None:
    ensure_directories()
    template_path = TMP_DIR / "sample_template.docx"
    build_template(template_path)

    rendered_docx = OUTPUT_DIR / "sample_rendered.docx"
    context = render_docx(template_path, rendered_docx)

    pdf_path = convert_with_libreoffice(rendered_docx, OUTPUT_DIR)
    html_path = convert_with_pandoc(rendered_docx, OUTPUT_DIR, "html")
    md_path = convert_with_pandoc(rendered_docx, OUTPUT_DIR, "md")

    summarize(
        {
            "context": context,
            "outputs": {
                "docx": str(rendered_docx),
                "pdf": str(pdf_path),
                "html": str(html_path),
                "markdown": str(md_path),
            },
        }
    )


if __name__ == "__main__":
    main()

